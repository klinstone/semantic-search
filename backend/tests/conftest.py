"""Общие фикстуры для тестов."""
import io

import pytest
from docx import Document as DocxDocument
from pypdf import PdfWriter


@pytest.fixture
def make_docx(tmp_path):
    """Фабрика DOCX-файлов с заданными параграфами и таблицами."""

    def _make(name: str, paragraphs: list[str], tables: list[list[list[str]]] | None = None):
        doc = DocxDocument()
        for p in paragraphs:
            doc.add_paragraph(p)
        for table_data in tables or []:
            rows = len(table_data)
            cols = len(table_data[0]) if rows > 0 else 0
            table = doc.add_table(rows=rows, cols=cols)
            for r, row in enumerate(table_data):
                for c, cell_text in enumerate(row):
                    table.cell(r, c).text = cell_text
        path = tmp_path / name
        doc.save(str(path))
        return path

    return _make


@pytest.fixture
def make_pdf(tmp_path):
    """Фабрика PDF-файлов из готовых текстовых страниц.

    Использует pypdf для генерации минимального валидного PDF.
    Текст добавляется как простой content stream — этого достаточно,
    чтобы pypdf на чтении его извлёк обратно.
    """
    from pypdf.generic import (
        ArrayObject,
        DecodedStreamObject,
        DictionaryObject,
        FloatObject,
        IndirectObject,
        NameObject,
        NumberObject,
        TextStringObject,
    )

    def _make(name: str, pages_text: list[str]):
        writer = PdfWriter()
        for text in pages_text:
            page = writer.add_blank_page(width=612, height=792)
            # Создаём content stream с текстовой командой PDF.
            # Tj — рисует текст; формат: BT /F1 12 Tf 50 750 Td (text) Tj ET.
            escaped = text.replace("\\", "\\\\").replace("(", "\\(").replace(")", "\\)")
            content = f"BT /F1 12 Tf 50 750 Td ({escaped}) Tj ET".encode("latin-1")
            stream = DecodedStreamObject()
            stream.set_data(content)
            page[NameObject("/Contents")] = stream

            # Минимальный шрифт.
            font = DictionaryObject({
                NameObject("/Type"): NameObject("/Font"),
                NameObject("/Subtype"): NameObject("/Type1"),
                NameObject("/BaseFont"): NameObject("/Helvetica"),
            })
            resources = DictionaryObject({
                NameObject("/Font"): DictionaryObject({NameObject("/F1"): font}),
            })
            page[NameObject("/Resources")] = resources

        path = tmp_path / name
        with path.open("wb") as f:
            writer.write(f)
        return path

    return _make